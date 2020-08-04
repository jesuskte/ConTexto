'''
Código, funciones y clases relacionadas a la carga y lectura de
diferentes tipos de archivo (word, txt, rtf, pdf inicialmente).
'''
from utils.auxiliares import verificar_crear_dir

# Clase lector


class Lector():
    def __init__(self, ubicacion_archivo):
        """

        :param ubicacion_archivo:
        """
        self.establecer_ubicacion(ubicacion_archivo)

    def establecer_ubicacion(self, ubicacion_archivo):
        """

        :param ubicacion_archivo:
        :return:
        """
        self.ubicacion_archivo = ubicacion_archivo

    def leer_txt(self, encoding="utf-8"):
        """

        :param encoding:
        :return:
        """
        salida = []
        with open(self.ubicacion_archivo, encoding=encoding) as fp:
            linea = fp.readline()
            while linea:
                try:
                    salida.append(linea.strip())
                    linea = fp.readline()
                except BaseException:
                    continue
        return '\n'.join(salida)

    def leer_word(self, extraer_medios, dir_medios):
        """

        :param extraer_medios:
        :param dir_medios:
        :return:
        """
        import docx2txt
        if extraer_medios is False:
            texto = docx2txt.process(self.ubicacion_archivo)
        else:
            verify_create_dir(dir_medios)
            texto = docx2txt.process(self.ubicacion_archivo, dir_medios)
        return texto

    def leer_pdf(self, por_paginas, ocr, preprocesamiento, lenguaje, oem, psm, password=None):
        def leer_pdf(self, por_paginas, ocr, preprocesamiento, lenguaje, oem, psm):
        """

        :param por_paginas:
        :param ocr:
        :param preprocesamiento:
        :param lenguaje:
        :param oem:
        :param psm:
        :return:
        """
        if ocr:
            from utils.ocr import OCR
            recog = OCR(preprocesamiento, lenguaje, oem, psm)
            paginas = recog.pdf_a_texto(self.ubicacion_archivo)
        else:
            try:
                from utils.auxiliares import leer_pdf_slate
                paginas = leer_pdf_slate(self.ubicacion_archivo, password)
            except:
                from utils.auxiliares import leer_pdf_pypdf
                paginas = leer_pdf_pypdf(self.ubicacion_archivo, password)
        # Se define la forma de retornar el texto
        if por_paginas:
            return paginas
        else:
            return ' '.join(paginas)

    def leer_rtf(self):
        """

        :return:
        """
        from utils.auxiliares import striprtf
        texto = []
        with open(self.ubicacion_archivo) as fp:
            linea = fp.readline()
            while linea:
                try:
                    texto.append(linea.strip())
                    linea = fp.readline()
                except BaseException:
                    continue
        texto = [striprtf(i) for i in texto]
        texto = ' '.join(texto)
        return texto

    def leer_imagen(self, preprocesamiento, lenguaje, oem, psm):
        """

        :param preprocesamiento:
        :param lenguaje:
        :param oem:
        :param psm:
        :return:
        """
        from utils.ocr import OCR
        recog = OCR(preprocesamiento, lenguaje, oem, psm)
        texto = recog.imagen_a_texto(self.ubicacion_archivo)
        return texto

    def archivo_a_texto(
            self,
            tipo='inferir',
            extraer_medios=False,
            dir_medios="temp/img_dir/",
            por_paginas=False,
            encoding="utf-8",
            ocr=False,
            preprocesamiento=3,
            lenguaje='spa',
            oem=2,
            psm=3,
            password=None):
        if tipo == 'inferir':
            tipo = self.ubicacion_archivo.split('.')[-1]
        if tipo in ['txt', 'csv']:
            return self.leer_txt(encoding)
        elif tipo == 'pdf':
            return self.leer_pdf(por_paginas, ocr, preprocesamiento, lenguaje, oem, psm, password)
        elif tipo == 'rtf':
            return self.leer_rtf()
        elif tipo in ['doc', 'docx']:
            return self.leer_word(extraer_medios, dir_medios)
        elif tipo in ['png', 'jpg', 'jpeg']:
            return self.leer_imagen(preprocesamiento, lenguaje, oem, psm)
        else:
            print(
                'Formato desconocido. Por favor ingrese un archivo en formato adecuado.')
            return None

# Clase escritor


class Escritor():
    def __init__(self, ubicacion_archivo, texto):
        """

        :param ubicacion_archivo:
        :param texto:
        """
        self.establecer_ubicacion(ubicacion_archivo)
        self.establecer_texto(texto)

    def establecer_ubicacion(self, ubicacion_archivo):
        """

        :param ubicacion_archivo:
        :return:
        """
        self.ubicacion_archivo = ubicacion_archivo

    def establecer_texto(self, texto):
        """

        :param texto:
        :return:
        """
        self.texto = texto

    def escribir_txt(self):
        """

        :return:
        """
        if isinstance(self.texto, list):
            self.texto = '\n\n|**|\n\n'.join(self.texto)
        # with open(self.ubicacion_archivo, 'w') as fp:
        with open(self.ubicacion_archivo, 'w', encoding="utf-8") as fp:
            fp.write(self.texto)

    def escribir_word(self):
        """

        :return:
        """
        from docx import Document
        documento = Document()
        if isinstance(self.texto, list):
            for i, page in enumerate(self.texto):
                documento.add_paragraph(page)
                if i < len(self.texto) - 1:
                    documento.add_page_break()
        else:
            documento.add_paragraph(self.texto)
        documento.save(self.ubicacion_archivo)

    def escribir_pdf(self):
        """

        :return:
        """
        import PyPDF2
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from io import BytesIO
        from textwrap import wrap

        def escribir_pagina(texto):
            """

            :param texto:
            :return:
            """
            texto = texto.split('\n')
            temp = BytesIO()
            can = canvas.Canvas(temp, pagesize=letter)
            t = can.beginText()
            t.setFont('Helvetica-Bold', 7)
            t.setCharSpace(0)
            t.setTextOrigin(50, 700)
            for linea in texto:
                sublineas = wrap(linea, 150)
                if len(sublineas) > 0:
                    for sublinea in sublineas:
                        t.textLine(sublinea)
                else:
                    t.textLine('')
            can.drawText(t)
            # can.drawString(5, 550, text)
            can.save()
            temp.seek(0)
            lector = PyPDF2.PdfFileReader(temp)
            return lector
        salida = PyPDF2.PdfFileWriter()
        if isinstance(self.texto, list):
            for pag in self.texto:
                lector = escribir_pagina(pag)
                salida.addPage(lector.getPage(0))
        else:
            lector = escribir_pagina(self.texto)
            salida.addPage(lector.getPage(0))
        with open(self.ubicacion_archivo, 'wb') as fp:
            salida.write(fp)

    def texto_a_archivo(self, tipo='inferir'):
        """

        :param tipo:
        :return:
        """
        if tipo == 'inferir':
            tipo = self.ubicacion_archivo.split('.')[-1]
        if tipo in ['txt', 'csv']:
            self.escribir_txt()
        elif tipo == 'pdf':
            self.escribir_pdf()
        elif tipo in ['doc', 'docx']:
            self.escribir_word()
        else:
            print('Formato desconocido. Se escribirá en un formato plano (.txt).')
            nueva_ruta = ''.join(self.ubicacion_archivo.split(
                '.')[:-1]) + '_{}.txt'.format(tipo)
            self.establecer_ubicacion(nueva_ruta)
            self.escribir_txt()


# Funciones que encapsulan el proceso de lectura y escritura de archivos
# de texto

def leer_texto(
        ubicacion_archivo,
        tipo='inferir',
        extraer_medios=False,
        dir_medios="temp/img_dir/",
        por_paginas=False,
        encoding="utf-8",
        ocr=False,
        preprocesamiento=3,
        lenguaje='spa',
        oem=2,
        psm=3, 
        password=None):
        psm=3):
    """

    :param ubicacion_archivo:
    :param tipo:
    :param extraer_medios:
    :param dir_medios:
    :param por_paginas:
    :param encoding:
    :param ocr:
    :param preprocesamiento:
    :param lenguaje:
    :param oem:
    :param psm:
    :return:
    """
    le = Lector(ubicacion_archivo)
    return le.archivo_a_texto(
        tipo,
        extraer_medios,
        dir_medios,
        por_paginas,
        encoding,
        ocr,
        preprocesamiento,
        lenguaje,
        oem,
        psm,
        password)


def escribir_texto(ubicacion_archivo, texto, tipo='inferir'):
    """

    :param ubicacion_archivo:
    :param texto:
    :param tipo:
    :return:
    """
    es = Escritor(ubicacion_archivo, texto)
    es.texto_a_archivo(tipo)
